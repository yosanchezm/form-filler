"""
Utilidades para manipulación de documentos DOCX.
"""

import re
from typing import Dict, List, Optional, Tuple
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

from config.settings import MIN_UNDERSCORES
from .text_utils import first_underscore_span, find_all_underscore_spans


# =========================
# Detección de tokens
# =========================
# Regex mejorado para capturar brackets incluyendo saltos de línea internos
BRACKET_RE = re.compile(r"\[[^\]]{1,200}\]", re.DOTALL)


def extract_bracket_tokens_from_doc(doc: Document) -> List[str]:
    """
    Extrae todos los tokens tipo [....] del documento.
    Maneja casos donde el token está partido entre varios runs.
    """
    tokens = set()
    
    for _, p in iter_all_paragraphs(doc):
        # Obtener texto completo del párrafo (uniendo todos los runs)
        txt = get_paragraph_full_text(p)
        
        # Buscar tokens bracket
        for m in BRACKET_RE.finditer(txt):
            token = m.group(0).strip()
            # Normalizar espacios internos
            token = re.sub(r'\s+', ' ', token)
            if len(token) > 3:  # Mínimo [x]
                tokens.add(token)
    
    return sorted(tokens)


def get_paragraph_full_text(paragraph) -> str:
    """
    Obtiene el texto completo de un párrafo, manejando runs fragmentados.
    """
    if paragraph.runs:
        return "".join(r.text for r in paragraph.runs)
    return paragraph.text or ""


# =========================
# Manipulación de párrafos
# =========================
def set_paragraph_text(paragraph, new_text: str) -> None:
    """
    Establece el texto de un párrafo, preservando el primer run.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_in_paragraph_all(paragraph, replacements: Dict[str, str]) -> bool:
    """
    Reemplaza múltiples tokens dentro de un párrafo (incluye placeholders [..]).
    Maneja tokens partidos entre runs escribiendo el texto completo en el primer run.
    
    Retorna True si se hizo algún reemplazo.
    """
    if not paragraph.runs:
        return False
    full = "".join(r.text for r in paragraph.runs)
    new = full
    for k, v in replacements.items():
        if v is None:
            continue
        new = new.replace(k, str(v))
    if new != full:
        set_paragraph_text(paragraph, new)
        return True
    return False


def fill_underscore_in_paragraph(paragraph, value: str, min_underscores: int = MIN_UNDERSCORES) -> bool:
    """
    Rellena la primera secuencia de underscores con el valor proporcionado.
    
    Retorna True si se hizo el reemplazo.
    """
    full = "".join(r.text for r in paragraph.runs)
    span = first_underscore_span(full, min_underscores)
    if not span:
        return False
    start, end = span
    line_len = end - start

    v = (str(value) if value is not None else "").strip()
    v_padded = v + (" " * max(0, line_len - len(v)))  # relleno visual

    new = full[:start] + v_padded + full[end:]
    set_paragraph_text(paragraph, new)
    return True


def fill_all_underscores_in_paragraph(
    paragraph, 
    values: List[str], 
    min_underscores: int = MIN_UNDERSCORES
) -> int:
    """
    Rellena TODAS las secuencias de underscores con los valores proporcionados.
    
    Útil para campos como: "a los ____ días del mes de ____ de 20__"
    
    Args:
        paragraph: Párrafo a modificar
        values: Lista de valores en orden
        min_underscores: Mínimo de underscores consecutivos
    
    Returns:
        Número de campos rellenados
    """
    full = "".join(r.text for r in paragraph.runs)
    spans = find_all_underscore_spans(full, min_underscores)
    
    if not spans or not values:
        return 0
    
    # Rellenar de atrás hacia adelante para no alterar índices
    filled = 0
    new_text = full
    for span, value in zip(reversed(spans), reversed(values[:len(spans)])):
        start, end = span
        line_len = end - start
        v = (str(value) if value is not None else "").strip()
        v_padded = v + (" " * max(0, line_len - len(v)))
        new_text = new_text[:start] + v_padded + new_text[end:]
        filled += 1
    
    set_paragraph_text(paragraph, new_text)
    return filled


# =========================
# Iteración de párrafos
# =========================
def iter_all_paragraphs(doc: Document):
    """
    Itera todos los párrafos del documento, incluyendo:
    - Párrafos del cuerpo principal
    - Párrafos en tablas
    - Párrafos en headers/footers
    - Párrafos en cuadros de texto (content controls)
    
    Yields: (identificador, paragraph)
    """
    # Párrafos del cuerpo principal
    for i, p in enumerate(doc.paragraphs):
        yield (f"paragraph:{i}", p)

    # Párrafos en tablas del cuerpo principal
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield (f"table:{ti} r:{ri} c:{ci} p:{pi}", p)
    
    # Párrafos en headers y footers
    for section_idx, section in enumerate(doc.sections):
        # Header
        if section.header:
            for pi, p in enumerate(section.header.paragraphs):
                yield (f"header:{section_idx} p:{pi}", p)
            # Tablas en header
            for ti, t in enumerate(section.header.tables):
                for ri, row in enumerate(t.rows):
                    for ci, cell in enumerate(row.cells):
                        for cpi, p in enumerate(cell.paragraphs):
                            yield (f"header:{section_idx} table:{ti} r:{ri} c:{ci} p:{cpi}", p)
        
        # Footer
        if section.footer:
            for pi, p in enumerate(section.footer.paragraphs):
                yield (f"footer:{section_idx} p:{pi}", p)
            # Tablas en footer
            for ti, t in enumerate(section.footer.tables):
                for ri, row in enumerate(t.rows):
                    for ci, cell in enumerate(row.cells):
                        for cpi, p in enumerate(cell.paragraphs):
                            yield (f"footer:{section_idx} table:{ti} r:{ri} c:{ci} p:{cpi}", p)


# =========================
# Detección de formato
# =========================
def is_run_bold(run) -> bool:
    """Verifica si un run está en negrita."""
    return run.bold is True


def is_run_underlined(run) -> bool:
    """Verifica si un run está subrayado."""
    return run.underline is not None and run.underline != WD_UNDERLINE.NONE


def is_run_highlighted(run) -> bool:
    """Verifica si un run tiene resaltado o color especial."""
    if run.font.highlight_color is not None:
        return True
    if run.font.color and run.font.color.rgb:
        # Colores comunes de "placeholder" (rojo, azul)
        rgb = run.font.color.rgb
        if rgb == RGBColor(255, 0, 0) or rgb == RGBColor(0, 0, 255):
            return True
    return False


def detect_formatted_placeholders(paragraph) -> List[Dict]:
    """
    Detecta texto con formato especial que podría ser un placeholder.
    Incluye detección de texto con fondo gris/resaltado.
    
    Returns:
        Lista de dicts con info de cada placeholder formateado
    """
    placeholders = []
    full_text = get_paragraph_full_text(paragraph)
    
    # Palabras clave comunes en formularios colombianos
    PLACEHOLDER_KEYWORDS = [
        "nombre", "incluir", "indicar", "razón social", "nit", "cédula",
        "fecha", "dirección", "teléfono", "correo", "email", "ciudad",
        "departamento", "país", "número", "proceso", "contratación",
        "representante", "legal", "revisor", "fiscal", "tarjeta", 
        "profesional", "firma", "identificación", "documento", "cargo",
        "objeto", "valor", "plazo", "domicilio", "constitucion", "vigencia"
    ]
    
    # Detectar runs con formato especial
    for i, run in enumerate(paragraph.runs):
        text = run.text.strip()
        if not text or len(text) < 2:
            continue
        
        formats = []
        if is_run_bold(run):
            formats.append("bold")
        if is_run_underlined(run):
            formats.append("underline")
        if is_run_highlighted(run):
            formats.append("highlight")
        if has_shading(run):
            formats.append("shading")
        
        # Determinar si parece un placeholder
        text_lower = text.lower()
        is_placeholder = (
            text.startswith("[") or 
            text.endswith("]") or
            text.isupper() or
            any(kw in text_lower for kw in PLACEHOLDER_KEYWORDS)
        )
        
        if formats and is_placeholder:
            placeholders.append({
                "text": text,
                "format": formats,
                "run_index": i
            })
    
    # También detectar texto resaltado con fondo gris (shading)
    # que puede ser un bloque completo
    shaded_text = extract_shaded_text(paragraph)
    for st in shaded_text:
        if st not in [p["text"] for p in placeholders]:
            placeholders.append({
                "text": st,
                "format": ["shading"],
                "run_index": -1
            })
    
    return placeholders


def has_shading(run) -> bool:
    """
    Verifica si un run tiene sombreado de fondo (shading).
    Común en documentos Word para resaltar campos a llenar.
    
    Detecta:
    - Shading con fill color (fondo gris, amarillo, etc.)
    - Shading con patternFill
    - Background colors en el run
    """
    try:
        run_xml = run._element
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # Buscar shading en las propiedades del run
        shading = run_xml.find(f'.//{ns}shd')
        if shading is not None:
            fill = shading.get(f'{ns}fill')
            color = shading.get(f'{ns}color')
            val = shading.get(f'{ns}val')
            
            # Si tiene fill y no es blanco ni transparente
            if fill and fill.lower() not in ('ffffff', 'auto', 'none', '', 'clear'):
                return True
            
            # Si tiene un patrón de shading
            if val and val.lower() not in ('nil', 'clear', 'none'):
                return True
            
            # Si tiene color de shading
            if color and color.lower() not in ('ffffff', 'auto', 'none', ''):
                return True
        
        # También verificar highlight (resaltado)
        highlight = run_xml.find(f'.//{ns}highlight')
        if highlight is not None:
            val = highlight.get(f'{ns}val')
            if val and val.lower() not in ('none', 'clear', 'white'):
                return True
                
    except Exception:
        pass
    return False


def extract_shaded_text(paragraph) -> List[str]:
    """
    Extrae texto que tiene fondo sombreado (shading) en el párrafo.
    """
    shaded_texts = []
    current_shaded = []
    
    for run in paragraph.runs:
        if has_shading(run) and run.text.strip():
            current_shaded.append(run.text)
        else:
            if current_shaded:
                text = "".join(current_shaded).strip()
                if text and len(text) > 3:
                    shaded_texts.append(text)
                current_shaded = []
    
    # No olvidar el último grupo
    if current_shaded:
        text = "".join(current_shaded).strip()
        if text and len(text) > 3:
            shaded_texts.append(text)
    
    return shaded_texts


# =========================
# Manipulación de tablas
# =========================
def get_table_structure(table) -> List[List[str]]:
    """
    Obtiene la estructura de una tabla como matriz de textos.
    
    Returns:
        Lista de filas, cada fila es lista de textos de celdas
    """
    structure = []
    for row in table.rows:
        row_texts = []
        for cell in row.cells:
            cell_text = " ".join(p.text for p in cell.paragraphs).strip()
            row_texts.append(cell_text)
        structure.append(row_texts)
    return structure


def find_label_value_pairs_in_table(table) -> List[Dict]:
    """
    Detecta pares label-valor en una tabla.
    
    Busca patrones comunes:
    - Columna 0 = label, Columna 1 = valor (vacío o underscores)
    - Fila header con labels y valores debajo
    
    Returns:
        Lista de dicts con info de cada par detectado
    """
    pairs = []
    
    for ri, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 2:
            continue
        
        label = " ".join(p.text for p in cells[0].paragraphs).strip()
        value = " ".join(p.text for p in cells[1].paragraphs).strip()
        
        if not label:
            continue
        
        # Detectar si el valor está vacío o tiene underscores
        is_empty = not value or all(c in "_ \t\n" for c in value)
        has_underscores = "_" * MIN_UNDERSCORES in value
        
        if is_empty or has_underscores:
            pairs.append({
                "row": ri,
                "label_col": 0,
                "value_col": 1,
                "label": label,
                "current_value": value,
                "is_empty": is_empty,
                "has_underscores": has_underscores
            })
    
    return pairs


# =========================
# Checkboxes
# =========================
def detect_checkbox_cells(table) -> List[Dict]:
    """
    Detecta celdas que parecen checkboxes en una tabla.
    
    Returns:
        Lista de checkboxes detectados con su posición
    """
    checkboxes = []
    
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = " ".join(p.text for p in cell.paragraphs).strip()
            
            # Patrones de checkbox
            if re.match(r"^[_\s]{2,}$", text):  # Solo underscores
                # Buscar label en celda adyacente o anterior
                label = ""
                if ci > 0:
                    label = " ".join(p.text for p in row.cells[ci-1].paragraphs).strip()
                
                checkboxes.append({
                    "row": ri,
                    "col": ci,
                    "label": label,
                    "type": "underscore_checkbox"
                })
            
            # Sí___ No___ pattern
            if re.search(r"(Sí|Si)\s*[_]+\s*(No)\s*[_]+", text, re.IGNORECASE):
                checkboxes.append({
                    "row": ri,
                    "col": ci,
                    "label": text,
                    "type": "yes_no"
                })
    
    return checkboxes


def fill_checkbox(cell, value: bool, style: str = "X") -> bool:
    """
    Marca o desmarca un checkbox en una celda.
    
    Args:
        cell: Celda de la tabla
        value: True para marcar, False para desmarcar
        style: Estilo de marca ("X", "✓", "●")
    
    Returns:
        True si se modificó
    """
    if not cell.paragraphs:
        return False
    
    text = " ".join(p.text for p in cell.paragraphs).strip()
    
    # Para checkbox Sí/No
    if "sí" in text.lower() or "no" in text.lower():
        new_text = text
        if value:
            # Marcar Sí
            new_text = re.sub(r"(Sí|Si)\s*[_]+", f"\\1 {style}", text, flags=re.IGNORECASE)
        else:
            # Marcar No
            new_text = re.sub(r"(No)\s*[_]+", f"\\1 {style}", text, flags=re.IGNORECASE)
        
        if new_text != text:
            set_paragraph_text(cell.paragraphs[0], new_text)
            return True
    
    # Para checkbox simple (solo underscores)
    if re.match(r"^[_\s]+$", text):
        mark = style if value else ""
        set_paragraph_text(cell.paragraphs[0], mark)
        return True
    
    return False


# =========================
# Utilidades generales
# =========================
def count_fillable_fields(doc: Document) -> Dict[str, int]:
    """
    Cuenta los campos rellenables en el documento por tipo.
    
    Returns:
        Dict con conteos por tipo de campo
    """
    counts = {
        "brackets": 0,
        "underscores": 0,
        "table_cells": 0,
        "checkboxes": 0,
        "formatted": 0,
        "total": 0
    }
    
    # Brackets
    tokens = extract_bracket_tokens_from_doc(doc)
    counts["brackets"] = len(tokens)
    
    # Underscores en párrafos
    for where, p in iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        if first_underscore_span(full, MIN_UNDERSCORES):
            if "table:" in where:
                counts["table_cells"] += 1
            else:
                counts["underscores"] += 1
        
        # Placeholders formateados
        formatted = detect_formatted_placeholders(p)
        counts["formatted"] += len(formatted)
    
    # Checkboxes en tablas
    for table in doc.tables:
        checkboxes = detect_checkbox_cells(table)
        counts["checkboxes"] += len(checkboxes)
    
    counts["total"] = sum(v for k, v in counts.items() if k != "total")
    return counts


def get_document_preview(doc: Document, max_paragraphs: int = 20) -> str:
    """
    Genera una vista previa del documento como texto.
    
    Args:
        doc: Documento DOCX
        max_paragraphs: Máximo de párrafos a incluir
    
    Returns:
        Texto con preview del documento
    """
    lines = []
    count = 0
    
    for where, p in iter_all_paragraphs(doc):
        if count >= max_paragraphs:
            lines.append("... [más contenido]")
            break
        
        text = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        if text.strip():
            lines.append(f"[{where}] {text[:200]}")
            count += 1
    
    return "\n".join(lines)

