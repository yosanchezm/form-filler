"""
Procesador principal del documento - Versión mejorada con LangChain.
"""

import io
from typing import Dict, List, Any, Tuple, Optional

from docx import Document

from models.events import FillEvent
from core.filler import (
    fill_bracket_placeholders,
    fill_paragraph_underscore_fields,
    fill_table_fields,
    fill_formatted_placeholders,
    apply_ai_updates,
    get_document_analysis
)
from ai.prompts import build_prompt
from ai.providers import call_llm

# Importar LangChain processor si está disponible
LANGCHAIN_AVAILABLE = False
try:
    from ai.langchain_processor import (
        call_llm_with_structured_output,
        is_langchain_supported,
        get_supported_providers
    )
    LANGCHAIN_AVAILABLE = True
except ImportError:
    pass


def run_mvp2(
    docx_bytes: bytes,
    kb: Dict[str, Dict[str, Any]],
    use_ai: bool,
    provider: str,
    api_key: str,
    model: str,
    confidence_threshold: float,
    temperature: float,
    use_langchain: bool = True,
    use_semantic_search: bool = True
) -> Tuple[bytes, List[FillEvent], List[FillEvent], List[Dict[str, Any]]]:
    """
    Procesa un documento DOCX rellenando campos con reglas y opcionalmente con IA.
    
    Args:
        docx_bytes: Bytes del documento DOCX
        kb: Base de conocimiento con claves '__raw__' y '__norm__'
        use_ai: Si se debe usar IA para completar campos pendientes
        provider: Proveedor de IA
        api_key: Clave API del proveedor
        model: Modelo a usar
        confidence_threshold: Umbral de confianza mínimo
        temperature: Temperatura del modelo
        use_langchain: Si se debe usar LangChain (default: True si disponible)
        use_semantic_search: Si se debe usar búsqueda semántica en KB
    
    Returns:
        - docx output bytes
        - filled_by_rules: Eventos rellenados por reglas
        - filled_by_ai: Eventos rellenados por IA
        - skipped_ai_updates: Actualizaciones de IA no aplicadas
    """
    doc = Document(io.BytesIO(docx_bytes))
    kb_norm = kb["__norm__"]
    kb_raw = kb["__raw__"]
    
    # KB completa para búsqueda híbrida
    kb_full = kb if use_semantic_search else None

    filled_rules: List[FillEvent] = []
    ai_targets: List[Dict[str, Any]] = []

    # 1) Bracket placeholders: [ ... ]
    fr, targets_br = fill_bracket_placeholders(doc, kb_norm, kb_full)
    filled_rules.extend(fr)
    ai_targets.extend(targets_br)

    # 2) Underscore fields in paragraphs
    fr, targets_u = fill_paragraph_underscore_fields(doc, kb_norm, kb_full)
    filled_rules.extend(fr)
    ai_targets.extend(targets_u)

    # 3) Table fields
    fr, targets_t = fill_table_fields(doc, kb_norm, kb_full)
    filled_rules.extend(fr)
    ai_targets.extend(targets_t)

    # 4) Formatted placeholders (negrita, subrayado, etc)
    fr, targets_f = fill_formatted_placeholders(doc, kb_norm, kb_full)
    filled_rules.extend(fr)
    ai_targets.extend(targets_f)

    filled_ai: List[FillEvent] = []
    skipped: List[Dict[str, Any]] = []

    # 5) AI completion for remaining targets
    if use_ai and ai_targets:
        llm_out = _call_ai(
            provider=provider,
            api_key=api_key,
            model=model,
            targets=ai_targets,
            kb_raw=kb_raw,
            temperature=temperature,
            use_langchain=use_langchain
        )

        updates = llm_out.get("updates", [])
        if not isinstance(updates, list):
            raise ValueError("La IA no devolvió 'updates' como lista.")

        # Agregar warnings si existen
        warnings = llm_out.get("warnings", [])
        if warnings:
            for w in warnings:
                skipped.append({"warning": w, "reason": "ai_warning"})

        applied, skipped_ai = apply_ai_updates(doc, updates, confidence_threshold)
        filled_ai.extend(applied)
        skipped.extend(skipped_ai)

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue(), filled_rules, filled_ai, skipped


def _call_ai(
    provider: str,
    api_key: str,
    model: str,
    targets: List[Dict[str, Any]],
    kb_raw: Dict[str, Any],
    temperature: float,
    use_langchain: bool = True
) -> Dict[str, Any]:
    """
    Llama a la IA usando LangChain o el método tradicional.
    
    Args:
        provider: Proveedor de IA
        api_key: Clave API
        model: Modelo a usar
        targets: Lista de campos a completar
        kb_raw: Base de conocimiento raw
        temperature: Temperatura del modelo
        use_langchain: Si usar LangChain
    
    Returns:
        Dict con 'updates' y opcionalmente 'warnings'
    """
    # Intentar usar LangChain si está disponible y el proveedor es soportado
    if use_langchain and LANGCHAIN_AVAILABLE and is_langchain_supported(provider):
        try:
            return call_llm_with_structured_output(
                provider=provider,
                api_key=api_key,
                model=model,
                targets=targets,
                knowledge_base=kb_raw,
                temperature=temperature
            )
        except Exception as e:
            # Fallback al método tradicional si falla LangChain
            print(f"LangChain falló, usando método tradicional: {e}")
    
    # Método tradicional
    system, user = build_prompt(targets, kb_raw)
    return call_llm(provider, api_key, model, system, user, temperature=temperature)


def analyze_document(docx_bytes: bytes) -> Dict[str, Any]:
    """
    Analiza un documento sin procesarlo, retornando estadísticas.
    
    Args:
        docx_bytes: Bytes del documento DOCX
    
    Returns:
        Dict con análisis del documento
    """
    doc = Document(io.BytesIO(docx_bytes))
    return get_document_analysis(doc)


def run_rules_only(
    docx_bytes: bytes,
    kb: Dict[str, Dict[str, Any]],
    use_semantic_search: bool = True
) -> Tuple[bytes, List[FillEvent], List[Dict[str, Any]]]:
    """
    Procesa el documento solo con reglas, sin IA.
    
    Útil para preview o cuando no se quiere usar IA.
    
    Args:
        docx_bytes: Bytes del documento DOCX
        kb: Base de conocimiento
        use_semantic_search: Si usar búsqueda semántica
    
    Returns:
        - docx output bytes
        - filled events
        - pending targets (que necesitarían IA)
    """
    doc = Document(io.BytesIO(docx_bytes))
    kb_norm = kb["__norm__"]
    kb_full = kb if use_semantic_search else None

    filled: List[FillEvent] = []
    pending: List[Dict[str, Any]] = []

    # Procesar con reglas
    fr, targets = fill_bracket_placeholders(doc, kb_norm, kb_full)
    filled.extend(fr)
    pending.extend(targets)

    fr, targets = fill_paragraph_underscore_fields(doc, kb_norm, kb_full)
    filled.extend(fr)
    pending.extend(targets)

    fr, targets = fill_table_fields(doc, kb_norm, kb_full)
    filled.extend(fr)
    pending.extend(targets)

    fr, targets = fill_formatted_placeholders(doc, kb_norm, kb_full)
    filled.extend(fr)
    pending.extend(targets)

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue(), filled, pending


def get_ai_availability() -> Dict[str, Any]:
    """
    Retorna información sobre la disponibilidad de funciones de IA.
    
    Returns:
        Dict con info de disponibilidad
    """
    result = {
        "langchain_available": LANGCHAIN_AVAILABLE,
        "traditional_available": True,
        "langchain_providers": []
    }
    
    if LANGCHAIN_AVAILABLE:
        result["langchain_providers"] = get_supported_providers()
    
    return result

