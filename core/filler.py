"""
Lógica de relleno de formularios DOCX - Versión mejorada.
"""

from typing import Dict, List, Any, Tuple, Optional

from docx import Document

from config.settings import MIN_UNDERSCORES
from models.events import FillEvent
from utils.docx_utils import (
    extract_bracket_tokens_from_doc,
    iter_all_paragraphs,
    replace_in_paragraph_all,
    fill_underscore_in_paragraph,
    fill_all_underscores_in_paragraph,
    set_paragraph_text,
    detect_formatted_placeholders,
    find_label_value_pairs_in_table,
    detect_checkbox_cells,
    fill_checkbox,
    count_fillable_fields,
    get_paragraph_full_text
)
from utils.text_utils import (
    first_underscore_span,
    find_all_underscore_spans,
    extract_label_before_underscores,
    looks_empty_or_underscores,
    detect_checkbox_field,
    detect_date_field,
    detect_highlighted_field,
    extract_field_context
)
from core.knowledge_base import find_value_for_label, find_value_hybrid


def fill_bracket_placeholders(
    doc: Document,
    kb_norm: Dict[str, str],
    kb_full: Optional[Dict] = None
) -> Tuple[List[FillEvent], List[Dict[str, Any]]]:
    """
    Detecta tokens [....] y trata de mapearlos a KB.
    
    Args:
        doc: Documento DOCX
        kb_norm: KB normalizada para búsqueda por reglas
        kb_full: KB completa para búsqueda híbrida (opcional)
    
    Returns:
        - filled events (rellenos por reglas)
        - targets para IA (si no se pudo resolver por reglas)
    """
    tokens = extract_bracket_tokens_from_doc(doc)
    filled: List[FillEvent] = []
    ai_targets: List[Dict[str, Any]] = []

    for tok in tokens:
        # Extraer contenido dentro del corchete como label
        inner = tok[1:-1].strip()
        
        # Intentar búsqueda híbrida si está disponible
        result = None
        if kb_full:
            result = find_value_hybrid(inner, kb_full)
        
        if result:
            val, confidence = result
            # Aplicar reemplazo global
            replaced_any = False
            for where, p in iter_all_paragraphs(doc):
                if replace_in_paragraph_all(p, {tok: str(val)}):
                    replaced_any = True
            if replaced_any:
                source = "rules" if confidence == 1.0 else "semantic"
                filled.append(FillEvent(
                    where="doc", 
                    label=tok, 
                    value=str(val), 
                    source=source, 
                    confidence=confidence
                ))
        else:
            # Fallback a búsqueda simple
            val = find_value_for_label(inner, kb_norm)
            if val is not None:
                replaced_any = False
                for where, p in iter_all_paragraphs(doc):
                    if replace_in_paragraph_all(p, {tok: str(val)}):
                        replaced_any = True
                if replaced_any:
                    filled.append(FillEvent(
                        where="doc", 
                        label=tok, 
                        value=str(val), 
                        source="rules", 
                        confidence=1.0
                    ))
            else:
                # Enviar a IA como target
                ai_targets.append({
                    "kind": "bracket",
                    "token": tok,
                    "label": inner
                })

    return filled, ai_targets


def fill_paragraph_underscore_fields(
    doc: Document,
    kb_norm: Dict[str, str],
    kb_full: Optional[Dict] = None
) -> Tuple[List[FillEvent], List[Dict[str, Any]]]:
    """
    Rellena campos con underscores en párrafos del documento.
    
    Returns:
        - filled events
        - targets para IA
    """
    filled: List[FillEvent] = []
    ai_targets: List[Dict[str, Any]] = []

    for i, p in enumerate(doc.paragraphs):
        full = "".join(r.text for r in p.runs)
        if not full:
            continue
        
        # Detectar si hay underscores
        spans = find_all_underscore_spans(full, MIN_UNDERSCORES)
        if not spans:
            continue
        
        # Extraer contexto del campo
        context = extract_field_context(full)
        label = context["label"]
        
        # Detectar tipos especiales de campos
        date_field = detect_date_field(full)
        checkbox = detect_checkbox_field(full)
        
        # Si es campo de fecha con múltiples espacios
        if date_field and len(spans) > 1:
            # Intentar obtener la fecha de la KB
            fecha_result = None
            if kb_full:
                # Buscar diferentes variantes de fecha
                for fecha_label in ["Fecha", "fecha", "Ciudad", "ciudad", "Fecha Firma"]:
                    fecha_result = find_value_hybrid(fecha_label, kb_full)
                    if fecha_result:
                        break
            
            if fecha_result:
                # Tenemos la fecha, enviar a IA para que la desglose
                ai_targets.append({
                    "kind": "date_field",
                    "where": f"paragraph:{i}",
                    "label": label if label else "Fecha de firma",
                    "sample": full[:300],
                    "num_fields": len(spans),
                    "hint": f"Desglosa esta fecha en los {len(spans)} espacios disponibles",
                    "available_value": fecha_result[0]
                })
            else:
                ai_targets.append({
                    "kind": "date_field",
                    "where": f"paragraph:{i}",
                    "label": label if label else "Fecha de firma",
                    "sample": full[:300],
                    "num_fields": len(spans),
                    "hint": f"Campo de fecha con {len(spans)} espacios (ej: ciudad, día, mes, año)"
                })
            continue
        
        # Si es checkbox, marcarlo como tal
        if checkbox:
            ai_targets.append({
                "kind": "checkbox",
                "where": f"paragraph:{i}",
                "label": label,
                "sample": full[:300],
                "checkbox_type": checkbox["type"],
                "options": checkbox.get("options", [])
            })
            continue
        
        # Buscar valor en KB
        result = None
        if kb_full:
            result = find_value_hybrid(label, kb_full)
        
        if result:
            val, confidence = result
            if fill_underscore_in_paragraph(p, str(val), MIN_UNDERSCORES):
                source = "rules" if confidence == 1.0 else "semantic"
                filled.append(FillEvent(
                    where=f"paragraph:{i}",
                    label=label,
                    value=str(val),
                    source=source,
                    confidence=confidence
                ))
        else:
            # Fallback
            val = find_value_for_label(label, kb_norm)
            if val is not None:
                if fill_underscore_in_paragraph(p, str(val), MIN_UNDERSCORES):
                    filled.append(FillEvent(
                        where=f"paragraph:{i}",
                        label=label,
                        value=str(val),
                        source="rules",
                        confidence=1.0
                    ))
            else:
                ai_targets.append({
                    "kind": "underscore",
                    "where": f"paragraph:{i}",
                    "label": label,
                    "sample": full[:300],
                    "context_before": context.get("before", ""),
                    "context_after": context.get("after", "")
                })
    
    return filled, ai_targets


def fill_table_fields(
    doc: Document,
    kb_norm: Dict[str, str],
    kb_full: Optional[Dict] = None
) -> Tuple[List[FillEvent], List[Dict[str, Any]]]:
    """
    Rellena campos en tablas del documento.
    
    Returns:
        - filled events
        - targets para IA
    """
    filled: List[FillEvent] = []
    ai_targets: List[Dict[str, Any]] = []

    for ti, table in enumerate(doc.tables):
        # Detectar checkboxes en la tabla
        checkboxes = detect_checkbox_cells(table)
        checkbox_positions = {(cb["row"], cb["col"]) for cb in checkboxes}
        
        for ri, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 2:
                continue

            label_text = (cells[0].text or "").strip()
            value_text = (cells[1].text or "").strip()

            if not label_text:
                continue

            # Verificar si es un checkbox
            if (ri, 1) in checkbox_positions:
                ai_targets.append({
                    "kind": "table_checkbox",
                    "where": f"table:{ti} row:{ri} col:1",
                    "label": label_text,
                    "sample": value_text[:200]
                })
                continue

            # Caso típico: label en col 0, valor en col 1
            if looks_empty_or_underscores(value_text):
                # Buscar valor con método híbrido si está disponible
                result = None
                if kb_full:
                    result = find_value_hybrid(label_text, kb_full)
                
                if result:
                    val, confidence = result
                    if not cells[1].paragraphs:
                        cells[1].add_paragraph(str(val))
                    else:
                        set_paragraph_text(cells[1].paragraphs[0], str(val))
                    
                    source = "rules" if confidence == 1.0 else "semantic"
                    filled.append(FillEvent(
                        where=f"table:{ti} row:{ri} col:1",
                        label=label_text,
                        value=str(val),
                        source=source,
                        confidence=confidence
                    ))
                else:
                    # Fallback
                    val = find_value_for_label(label_text, kb_norm)
                    if val is not None:
                        if not cells[1].paragraphs:
                            cells[1].add_paragraph(str(val))
                        else:
                            set_paragraph_text(cells[1].paragraphs[0], str(val))
                        filled.append(FillEvent(
                            where=f"table:{ti} row:{ri} col:1",
                            label=label_text,
                            value=str(val),
                            source="rules",
                            confidence=1.0
                        ))
                    else:
                        ai_targets.append({
                            "kind": "table_cell",
                            "where": f"table:{ti} row:{ri} col:1",
                            "label": label_text,
                            "sample": (cells[0].text or "")[:200]
                        })
            else:
                # Si el valor está dentro de un párrafo con underscores
                for pi, p in enumerate(cells[1].paragraphs):
                    full = "".join(r.text for r in p.runs)
                    if first_underscore_span(full, MIN_UNDERSCORES):
                        result = None
                        if kb_full:
                            result = find_value_hybrid(label_text, kb_full)
                        
                        if result:
                            val, confidence = result
                            if fill_underscore_in_paragraph(p, str(val), MIN_UNDERSCORES):
                                source = "rules" if confidence == 1.0 else "semantic"
                                filled.append(FillEvent(
                                    where=f"table:{ti} row:{ri} col:1 p:{pi}",
                                    label=label_text,
                                    value=str(val),
                                    source=source,
                                    confidence=confidence
                                ))
                        else:
                            val = find_value_for_label(label_text, kb_norm)
                            if val is not None:
                                if fill_underscore_in_paragraph(p, str(val), MIN_UNDERSCORES):
                                    filled.append(FillEvent(
                                        where=f"table:{ti} row:{ri} col:1 p:{pi}",
                                        label=label_text,
                                        value=str(val),
                                        source="rules",
                                        confidence=1.0
                                    ))
                            else:
                                ai_targets.append({
                                    "kind": "table_cell_underscore",
                                    "where": f"table:{ti} row:{ri} col:1 p:{pi}",
                                    "label": label_text,
                                    "sample": full[:300]
                                })
                        break

    return filled, ai_targets


def fill_formatted_placeholders(
    doc: Document,
    kb_norm: Dict[str, str],
    kb_full: Optional[Dict] = None
) -> Tuple[List[FillEvent], List[Dict[str, Any]]]:
    """
    Detecta y rellena placeholders con formato especial (negrita, subrayado, fondo gris, etc).
    
    Returns:
        - filled events
        - targets para IA
    """
    filled: List[FillEvent] = []
    ai_targets: List[Dict[str, Any]] = []
    
    # Track de textos ya procesados para evitar duplicados
    processed_texts = set()

    for where, p in iter_all_paragraphs(doc):
        placeholders = detect_formatted_placeholders(p)
        
        for ph in placeholders:
            text = ph["text"].strip()
            
            # Evitar duplicados
            if text in processed_texts:
                continue
            processed_texts.add(text)
            
            # Si es un bracket placeholder, se maneja en otra función
            if text.startswith("[") and text.endswith("]"):
                continue
            
            # Limpiar el texto para búsqueda
            search_text = text
            # Quitar prefijos comunes
            for prefix in ["Incluir ", "Indicar ", "Escribir "]:
                if search_text.lower().startswith(prefix.lower()):
                    search_text = search_text[len(prefix):]
            
            # Intentar encontrar valor
            result = None
            if kb_full:
                result = find_value_hybrid(search_text, kb_full)
            
            if result:
                val, confidence = result
                # Reemplazar el texto formateado
                if replace_in_paragraph_all(p, {text: str(val)}):
                    source = "rules" if confidence == 1.0 else "semantic"
                    filled.append(FillEvent(
                        where=where,
                        label=text,
                        value=str(val),
                        source=source,
                        confidence=confidence
                    ))
            else:
                val = find_value_for_label(search_text, kb_norm)
                if val is not None:
                    if replace_in_paragraph_all(p, {text: str(val)}):
                        filled.append(FillEvent(
                            where=where,
                            label=text,
                            value=str(val),
                            source="rules",
                            confidence=1.0
                        ))
                else:
                    # Agregar como target para IA con más contexto
                    full_text = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
                    ai_targets.append({
                        "kind": "highlighted",
                        "where": where,
                        "label": text,
                        "format": ph["format"],
                        "sample": full_text[:300],
                        "hint": f"Texto con formato especial ({', '.join(ph['format'])})"
                    })
    
    return filled, ai_targets


def apply_ai_updates(
    doc: Document,
    updates: List[Dict[str, Any]],
    confidence_threshold: float
) -> Tuple[List[FillEvent], List[Dict[str, Any]]]:
    """
    Aplica las actualizaciones sugeridas por IA al documento.
    
    Args:
        doc: Documento DOCX
        updates: Lista de actualizaciones de la IA
        confidence_threshold: Umbral mínimo de confianza para aplicar
    
    Returns:
        - applied events
        - skipped updates (no aplicados)
    """
    applied: List[FillEvent] = []
    skipped: List[Dict[str, Any]] = []

    for upd in updates:
        conf = float(upd.get("confidence", 0))
        if conf < confidence_threshold:
            skipped.append({**upd, "reason": f"confidence<{confidence_threshold}"})
            continue

        kind = upd.get("kind")
        value = upd.get("value")
        
        if value is None:
            skipped.append({**upd, "reason": "value is null"})
            continue

        success = False
        
        if kind == "bracket":
            success = _apply_bracket_update(doc, upd, applied)
        
        elif kind in ("underscore", "table_cell_underscore", "date_field"):
            success = _apply_underscore_update(doc, upd, applied)
        
        elif kind == "table_cell":
            success = _apply_table_cell_update(doc, upd, applied)
        
        elif kind in ("checkbox", "table_checkbox"):
            success = _apply_checkbox_update(doc, upd, applied)
        
        elif kind == "highlighted":
            success = _apply_highlighted_update(doc, upd, applied)
        
        else:
            skipped.append({**upd, "reason": f"unknown kind '{kind}'"})
            continue
        
        if not success:
            skipped.append({**upd, "reason": f"could not locate {kind} field to apply"})

    return applied, skipped


def _apply_bracket_update(
    doc: Document,
    upd: Dict[str, Any],
    applied: List[FillEvent]
) -> bool:
    """Aplica actualización de tipo bracket."""
    token = upd.get("token")
    value = upd.get("value")
    conf = float(upd.get("confidence", 0))
    
    if not token:
        return False
    
    replaced_any = False
    for where, p in iter_all_paragraphs(doc):
        if replace_in_paragraph_all(p, {token: str(value)}):
            replaced_any = True
    
    if replaced_any:
        applied.append(FillEvent(
            where="doc",
            label=token,
            value=str(value),
            source="ai",
            confidence=conf
        ))
        return True
    
    return False


def _apply_underscore_update(
    doc: Document,
    upd: Dict[str, Any],
    applied: List[FillEvent]
) -> bool:
    """Aplica actualización de tipo underscore o date_field."""
    label = (upd.get("label") or "").strip()
    value = upd.get("value")
    conf = float(upd.get("confidence", 0))
    kind = upd.get("kind", "underscore")
    
    # Para campos de fecha, el value puede venir como lista o string
    # Ej: value = "Bogotá, 15 de enero de 2025" o ["Bogotá", "15", "enero", "2025"]
    
    for where, p in iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        
        # Verificar si el label está en el párrafo
        label_matches = (not label) or (label in full)
        has_underscores = first_underscore_span(full, MIN_UNDERSCORES)
        
        if label_matches and has_underscores:
            # Caso especial: campo de fecha con múltiples espacios
            if kind == "date_field" or detect_date_field(full):
                spans = find_all_underscore_spans(full, MIN_UNDERSCORES)
                
                if len(spans) > 1:
                    # Si el value es una lista, usarla directamente
                    if isinstance(value, list):
                        values = value
                    else:
                        # Intentar parsear la fecha
                        values = _parse_date_value(str(value), len(spans))
                    
                    # Rellenar cada campo
                    if fill_all_underscores_in_paragraph(p, values, MIN_UNDERSCORES):
                        applied.append(FillEvent(
                            where=where,
                            label=label or "fecha",
                            value=str(value),
                            source="ai",
                            confidence=conf
                        ))
                        return True
            else:
                # Campo simple de underscore
                if fill_underscore_in_paragraph(p, str(value), MIN_UNDERSCORES):
                    applied.append(FillEvent(
                        where=where,
                        label=label,
                        value=str(value),
                        source="ai",
                        confidence=conf
                    ))
                    return True
    
    return False


def _parse_date_value(value: str, num_fields: int) -> List[str]:
    """
    Parsea un valor de fecha en componentes para rellenar campos individuales.
    
    Args:
        value: Valor de fecha (ej: "Bogotá, 15 de enero de 2025")
        num_fields: Número de campos underscore disponibles
    
    Returns:
        Lista de valores para cada campo
    """
    import re
    
    # Patrones comunes en fechas colombianas
    # "Bogotá, 15 de enero de 2025"
    # "15 de enero de 2025"
    # "enero 15, 2025"
    
    # Intentar extraer ciudad, día, mes, año
    city = ""
    day = ""
    month = ""
    year = ""
    
    # Buscar ciudad (antes de la coma)
    if "," in value:
        parts = value.split(",", 1)
        city = parts[0].strip()
        rest = parts[1].strip()
    else:
        rest = value.strip()
    
    # Buscar año (4 dígitos)
    year_match = re.search(r'\b(20\d{2}|19\d{2})\b', rest)
    if year_match:
        year = year_match.group(1)
    
    # Buscar día (1-31)
    day_match = re.search(r'\b(\d{1,2})\b(?!\d)', rest)
    if day_match:
        day = day_match.group(1)
    
    # Buscar mes (nombres en español)
    months_es = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    for m in months_es:
        if m in rest.lower():
            month = m.capitalize()
            break
    
    # Construir lista según número de campos
    if num_fields == 4:
        return [city or "", day, month, year]
    elif num_fields == 3:
        # Día, mes, año
        return [day, month, year]
    elif num_fields == 2:
        # Ciudad/día y mes/año
        if city:
            return [f"{city}, {day}", f"{month} de {year}"]
        return [day, f"{month} de {year}"]
    else:
        # Todo junto
        return [value]


def _apply_table_cell_update(
    doc: Document,
    upd: Dict[str, Any],
    applied: List[FillEvent]
) -> bool:
    """Aplica actualización de tipo table_cell."""
    label = (upd.get("label") or "").strip()
    value = upd.get("value")
    conf = float(upd.get("confidence", 0))
    
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            if len(row.cells) < 2:
                continue
            
            cell_label = (row.cells[0].text or "").strip()
            if cell_label == label:
                if row.cells[1].paragraphs:
                    set_paragraph_text(row.cells[1].paragraphs[0], str(value))
                else:
                    row.cells[1].add_paragraph(str(value))
                
                applied.append(FillEvent(
                    where=f"table:{ti} row:{ri} col:1",
                    label=label,
                    value=str(value),
                    source="ai",
                    confidence=conf
                ))
                return True
    
    return False


def _apply_checkbox_update(
    doc: Document,
    upd: Dict[str, Any],
    applied: List[FillEvent]
) -> bool:
    """Aplica actualización de tipo checkbox."""
    label = (upd.get("label") or "").strip()
    value = upd.get("value")
    conf = float(upd.get("confidence", 0))
    
    # El valor puede ser "Sí", "No", "X", True, False, etc.
    is_checked = str(value).lower() in ("sí", "si", "yes", "x", "true", "1", "✓")
    
    # Buscar en tablas primero
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = (cell.text or "").strip()
                if label in cell_text or cell_text in label:
                    if fill_checkbox(cell, is_checked):
                        applied.append(FillEvent(
                            where=f"table:{ti} row:{ri} col:{ci}",
                            label=label,
                            value=str(value),
                            source="ai",
                            confidence=conf
                        ))
                        return True
    
    # Buscar en párrafos
    for where, p in iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        if label and label in full:
            # Intentar marcar checkbox en el párrafo
            checkbox = detect_checkbox_field(full)
            if checkbox:
                # Modificar según el tipo
                new_text = full
                if checkbox["type"] == "yes_no":
                    if is_checked:
                        new_text = full.replace("Sí___", "Sí X").replace("Si___", "Si X")
                    else:
                        new_text = full.replace("No___", "No X")
                
                if new_text != full:
                    set_paragraph_text(p, new_text)
                    applied.append(FillEvent(
                        where=where,
                        label=label,
                        value=str(value),
                        source="ai",
                        confidence=conf
                    ))
                    return True
    
    return False


def _apply_highlighted_update(
    doc: Document,
    upd: Dict[str, Any],
    applied: List[FillEvent]
) -> bool:
    """Aplica actualización de tipo highlighted."""
    label = (upd.get("label") or "").strip()
    value = upd.get("value")
    conf = float(upd.get("confidence", 0))
    
    for where, p in iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        if label and label in full:
            if replace_in_paragraph_all(p, {label: str(value)}):
                applied.append(FillEvent(
                    where=where,
                    label=label,
                    value=str(value),
                    source="ai",
                    confidence=conf
                ))
                return True
    
    return False


def get_document_analysis(doc: Document) -> Dict[str, Any]:
    """
    Analiza un documento y retorna estadísticas de campos detectados.
    
    Returns:
        Dict con análisis del documento
    """
    counts = count_fillable_fields(doc)
    
    # Extraer todos los tokens bracket
    brackets = extract_bracket_tokens_from_doc(doc)
    
    # Analizar campos de fecha
    date_fields = []
    # Campos formateados (negrita, sombreado, etc)
    formatted_fields = []
    
    for where, p in iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        
        # Detectar campos de fecha
        if detect_date_field(full):
            date_fields.append({"where": where, "sample": full[:100]})
        
        # Detectar campos formateados
        placeholders = detect_formatted_placeholders(p)
        for ph in placeholders:
            formatted_fields.append({
                "where": where, 
                "text": ph["text"][:80],
                "format": ph["format"]
            })
    
    return {
        "counts": counts,
        "bracket_tokens": brackets,
        "date_fields": date_fields,
        "formatted_fields": formatted_fields[:50],  # Limitar para no sobrecargar
        "total_tables": len(doc.tables),
        "total_paragraphs": len(doc.paragraphs)
    }
